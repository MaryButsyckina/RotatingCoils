import Libs
import os
import openpyxl as excel

import docx as wr

from datetime import datetime
from Mathematics import CalculationsForExcel
from openpyxl.styles import PatternFill, Alignment
from openpyxl.chart import Reference, ScatterChart, Series
from openpyxl.chart.axis import ChartLines


class GraphTemplate:
    def __init__(self, selected_meas, a_coef_volts, b_coef_volts, comments, is_sextuple):
        super(GraphTemplate, self).__init__()

        self.book = None  # excel book, scroll to create_file
        self.file_name = None  # name of excel file
        self.analysis_folder_path = None
        self.graph_sheet = None  # active excel sheet
        self.a_coef_volts = a_coef_volts
        self.b_coef_volts = b_coef_volts
        self.selected_meas = selected_meas
        self.counter_of_I = None

        self.main_harm = 2
        if is_sextuple:
            self.main_harm = 3

        self.get_file()
        self.create_graph_sheet()
        self.write_data(comments)

    def get_file(self):
        full_path = None
        folder_path = ''

        tree = Libs.ET.parse('Configuration.cfg')
        root = tree.getroot()
        for elem in root[0].iter('PathOfFile'):
            full_path = elem.text

        split_path = full_path.split('/')
        self.file_name = split_path[-1].split('.')[0]
        for i in range(len(split_path) - 1):
            if i == len(split_path) - 2:
                folder_path += f'{split_path[i]}'
            else:
                folder_path += f'{split_path[i]}//'

        self.analysis_folder_path = f'{folder_path}//Analysis_of_{self.file_name}'

        if os.path.exists(self.analysis_folder_path):
            pass
        else:
            os.mkdir(self.analysis_folder_path)

        try:
            self.book = excel.load_workbook(f'{self.analysis_folder_path}//{self.file_name}.xlsx')
        except FileNotFoundError:
            self.book = excel.Workbook()
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        self.create_graph_sheet()

    def create_graph_sheet(self):
        date = ''
        date_array = f'{datetime.today()}'.split(' ')[0].split('-')
        for i in range(1, 4):
            if i == 3:
                date += f'{date_array[-i][2:]}'
            else:
                date += f'{date_array[-i]}.'

        try:
            self.graph_sheet = self.book[f'Graph_{date}']
        except KeyError:
            self.book.create_sheet(f'Graph_{date}', 0)
            self.graph_sheet = self.book[f'Graph_{date}']
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        try:
            del self.book['Sheet']
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')
        except KeyError:
            pass

    @staticmethod
    def get_data(a_coef_volts, b_coef_volts, selected_meas, main_harm):
        cls = CalculationsForExcel(a_coef_volts, b_coef_volts, selected_meas, main_harm)
        alpha = cls.calculate_alpha()
        coef = cls.calculate_coefficients_for_graph(alpha)

        return [alpha, coef]

    def write_data(self, comments):
        gd = self.get_data(self.a_coef_volts, self.b_coef_volts, self.selected_meas, self.main_harm)
        alpha = gd[0]
        a_coef_volts = gd[1][0]
        b_coef_volts = gd[1][1]

        self.graph_sheet.cell(row=1, column=3).value = alpha

        x = 0
        current_counter = 0
        for h in range(15):
            current_counter = 1
            y = 1
            if h > 0:
                x += 2

            self.graph_sheet.cell(row=2 + x, column=3).value = a_coef_volts[h][0]
            self.graph_sheet.cell(row=2 + x, column=2).value = comments[0]
            for n in range(1, len(self.selected_meas)):
                if comments[n] == comments[n - 1]:
                    self.graph_sheet.cell(row=2 + x, column=3 + y).value = a_coef_volts[h][n]
                    y += 1
                else:
                    x += 1
                    y = 1
                    current_counter += 1
                    self.graph_sheet.cell(row=2 + x, column=2).value = comments[n]
                    self.graph_sheet.cell(row=2 + x, column=3).value = a_coef_volts[h][n]

            x += 2
            y = 1
            self.graph_sheet.cell(row=2 + x, column=3).value = b_coef_volts[h][0]
            self.graph_sheet.cell(row=2 + x, column=2).value = comments[0]
            for n in range(1, len(self.selected_meas)):
                if comments[n] == comments[n - 1]:
                    self.graph_sheet.cell(row=2 + x, column=3 + y).value = b_coef_volts[h][n]
                    y += 1
                else:
                    x += 1
                    y = 1
                    self.graph_sheet.cell(row=2 + x, column=2).value = comments[n]
                    self.graph_sheet.cell(row=2 + x, column=3).value = b_coef_volts[h][n]

        rows_counter = x + 2

        self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        self.style_sheet(rows_counter, current_counter)
        self.add_graph(current_counter)

    def style_sheet(self, rows_counter, current_counter):
        self.graph_sheet.cell(row=1, column=2).value = 'angle'
        for rows in self.graph_sheet.iter_rows(min_row=1, max_row=1, min_col=2, max_col=3):
            for cell in rows:
                cell.fill = PatternFill(start_color='00FFCC00', end_color='00FFCC00', fill_type='solid')

        harm_counter = 0
        for i in range(2, rows_counter + 1, 2 * current_counter + 2):
            harm_counter += 1

            self.graph_sheet.merge_cells(f'A{i}:A{i + current_counter - 1}')
            self.graph_sheet.cell(row=i, column=1).value = f'a{harm_counter}'
            self.graph_sheet.cell(row=i, column=1).fill = PatternFill(start_color='0099CCFF', end_color='0099CCFF',
                                                                      fill_type='solid')
            self.graph_sheet.cell(row=i, column=1).alignment = Alignment(horizontal='center', vertical='center')

            self.graph_sheet.merge_cells(f'A{i + current_counter + 1}:A{i + 2 * current_counter}')
            self.graph_sheet.cell(row=i + current_counter + 1, column=1).value = f'b{harm_counter}'
            self.graph_sheet.cell(row=i + current_counter + 1, column=1).fill = PatternFill(start_color='00FF8080',
                                                                                            end_color='00FF8080',
                                                                                            fill_type='solid')
            self.graph_sheet.cell(row=i + current_counter + 1, column=1).alignment = Alignment(horizontal='center', vertical='center')

        self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

    def add_graph(self, current_counter):
        chart = []
        column_number = 65
        column_letter = chr(column_number)
        row_number = 1
        meas_count = -1

        for h in range(15):
            meas_count += 1
            chart.append(ScatterChart())
            chart[meas_count].width = 17
            chart[meas_count].height = 10
            chart[meas_count].y_axis.title = f'a{h+1}, unit'
            chart[meas_count].x_axis.tickLblPos = 'low'
            chart[meas_count].x_axis.title = 'current, A'
            chart[meas_count].x_axis.minorGridlines = ChartLines()
            xvalues = Reference(self.graph_sheet, min_col=2, min_row=2, max_row=1+current_counter)
            for i in range(30):
                values = Reference(self.graph_sheet, min_row=2+meas_count*(current_counter + 1), min_col=3+i, max_col=3+i, max_row=2+meas_count*(current_counter + 1) + current_counter - 1)
                series = Series(values, xvalues)
                series.marker.symbol = 'square'
                series.marker.size = 8
                series.marker.graphicalProperties.line.solidFill = 'FF0000'
                series.graphicalProperties.line.noFill = True
                chart[meas_count].series.append(series)

            if h == 3 or h == 6 or h == 9 or h == 12:
                column_number += 10
                if len(column_letter) > 1:
                    column_letter = f'A{chr(column_number)}'
                else:
                    column_letter = chr(column_number)
                if 90 < column_number:
                    column_number -= 25
                    column_letter = f'A{chr(column_number)}'

            if h == 0 or h == 3 or h == 6 or h == 9 or h == 12:
                row_number = 1
            else:
                row_number += 19

            self.graph_sheet.add_chart(chart[meas_count], f'{column_letter}{row_number}')

            meas_count += 1
            chart.append(ScatterChart())
            chart[meas_count].width = 17
            chart[meas_count].height = 10
            chart[meas_count].y_axis.title = f'b{h+1}, unit'
            chart[meas_count].x_axis.tickLblPos = 'low'
            chart[meas_count].x_axis.title = 'current, A'
            chart[meas_count].x_axis.minorGridlines = ChartLines()
            xvalues = Reference(self.graph_sheet, min_col=2, min_row=2, max_row=1+current_counter)
            for i in range(30):
                values = Reference(self.graph_sheet, min_row=2+meas_count*(current_counter+1), min_col=3 + i,
                                   max_col=3 + i, max_row=2+meas_count*(current_counter+1) + current_counter - 1)
                series = Series(values, xvalues)
                series.marker.symbol = 'square'
                series.marker.size = 8
                series.marker.graphicalProperties.line.solidFill = 'FF0000'
                series.graphicalProperties.line.noFill = True
                chart[meas_count].series.append(series)

            row_number += 19

            self.graph_sheet.add_chart(chart[meas_count], f'{column_letter}{row_number}')

        self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')


class MeasTemplate:
    def __init__(self, selected_meas, a_coef, b_coef, field, comments, is_sextuple):
        super(MeasTemplate, self).__init__()

        self.file_name = None
        self.analysis_folder_path = None
        self.book = None
        self.meas_sheet = None

        self.a_coef = a_coef
        self.b_coef = b_coef
        self.field = field
        self.selected_meas = selected_meas
        self.comments = comments
        self.relative_field = [[] for y in range(len(selected_meas))]

        self.main_harm = 1
        if is_sextuple:
            self.main_harm = 2

        self.get_file()
        self.get_data()
        self.write_data()


    def get_file(self):
        full_path = None
        folder_path = ''

        tree = Libs.ET.parse('Configuration.cfg')
        root = tree.getroot()
        for elem in root[0].iter('PathOfFile'):
            full_path = elem.text

        split_path = full_path.split('/')
        self.file_name = split_path[-1].split('.')[0]
        for i in range(len(split_path) - 1):
            if i == len(split_path) - 2:
                folder_path += f'{split_path[i]}'
            else:
                folder_path += f'{split_path[i]}//'

        self.analysis_folder_path = f'{folder_path}//Analysis_of_{self.file_name}'

        if os.path.exists(self.analysis_folder_path):
            pass
        else:
            os.mkdir(self.analysis_folder_path)

        try:
            self.book = excel.load_workbook(f'{self.analysis_folder_path}//{self.file_name}.xlsx')
        except FileNotFoundError:
            self.book = excel.Workbook()
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        self.create_meas_sheet()


    def create_meas_sheet(self):
        date = ''
        date_array = f'{datetime.today()}'.split(' ')[0].split('-')

        for i in range(1, 4):
            if i == 3:
                date += f'{date_array[-i][2:]}'
            else:
                date += f'{date_array[-i]}.'

        try:
            self.meas_sheet = self.book[f'Meas_{date}']
        except KeyError:
            self.book.create_sheet(f'Meas_{date}')
            self.meas_sheet = self.book[f'Meas_{date}']
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        try:
            del self.book['Sheet']
            self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')
        except KeyError:
            pass


    def get_data(self):
        i = 0
        for n in self.selected_meas:
            self.relative_field.append([])
            for h in range(15):
                self.relative_field[i].append(self.field[n][h]/self.field[n][self.main_harm])
            i += 1


    def write_data(self):
        counter = 3
        i = 0
        for n in self.selected_meas:
            self.meas_sheet.cell(row=2, column=counter).value = self.comments[n]
            for h in range(15):
                self.meas_sheet.cell(row=3 + h, column=counter).value = self.a_coef[n][h]
                self.meas_sheet.cell(row=19 + h, column=counter).value = self.b_coef[n][h]
                self.meas_sheet.cell(row=35 + h, column=counter).value = self.field[n][h]
                self.meas_sheet.cell(row=51 + h, column=counter).value = self.relative_field[i][h]

            counter += 1
            i += 1

        self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')

        self.style_sheet()


    def style_sheet(self):
        self.meas_sheet.column_dimensions['A'].width = 34
        self.meas_sheet.cell(row=2, column=2).value = 'current [A]'

        self.meas_sheet.merge_cells('A3:A17')
        self.meas_sheet.cell(row=3, column=1).value = 'A Gs'
        self.meas_sheet.cell(row=3, column=1).fill = PatternFill(start_color='0099CCFF', end_color='0099CCFF', fill_type='solid')
        self.meas_sheet.cell(row=3, column=1).alignment = Alignment(horizontal='center', vertical='center')

        self.meas_sheet.merge_cells('A19:A33')
        self.meas_sheet.cell(row=19, column=1).value = 'B Gs'
        self.meas_sheet.cell(row=19, column=1).fill = PatternFill(start_color='00FF8080', end_color='00FF8080', fill_type='solid')
        self.meas_sheet.cell(row=19, column=1).alignment = Alignment(horizontal='center', vertical='center')

        self.meas_sheet.merge_cells('A35:A49')
        self.meas_sheet.cell(row=35, column=1).value = 'Amplitude Gs'
        self.meas_sheet.cell(row=35, column=1).alignment = Alignment(horizontal='center', vertical='center')

        self.meas_sheet.merge_cells('A51:A65')
        self.meas_sheet.cell(row=51, column=1).value = 'Relative'
        self.meas_sheet.cell(row=51, column=1).alignment = Alignment(horizontal='center', vertical='center')

        for i in range(4):
            for h in range(15):
                self.meas_sheet.cell(row=3 + h + i*16, column=2).value = f'{h+1}'

        self.book.save(f'{self.analysis_folder_path}//{self.file_name}.xlsx')


class WordUtkinSextuple:
    def __init__(self, data):
        file = wr.Document('C:\\Users\\LAdmin\\PycharmProjects\\RotatingCoilsFinal\\Utkin_Sextuple_Template.docx')
        self.book_mark = list(filter(lambda paragraph: paragraph.text == '9. Magnetic measurements', file.paragraphs))[0]

        self.table_param = file.add_table(6, 2)
        self.table_meas = file.add_table(15, 5)
        self.table_param.style = 'Table Grid'
        self.table_meas.style = 'Table Grid'
        self.empty_paragraph = file.add_paragraph('')

        self.book_mark._p.addnext(self.table_meas._tbl)
        self.book_mark._p.addnext(self.empty_paragraph._p)
        self.book_mark._p.addnext(self.table_param._tbl)

        self.data = data

        self.write_param_data()
        self.write_meas_data()
        file.save('C:\\Users\\LAdmin\\Desktop\\Measurements\\Skif\\Sextuple\\6\\Analysis_of_Sextuple_6\\Hello.docx')

        self.table_param.cell(3, 0).add_paragraph(f'X(cm)')
        self.table_param.cell(4, 0).add_paragraph(f'Y(cm)')
        self.table_param.cell(5, 0).add_paragraph(f'Angle(rad)')

        self.table_param.cell(0, 1).add_paragraph(f'{data["Lens_length"]}')
        self.table_param.cell(1, 1).add_paragraph(f'{data["Radius"]}')
        self.table_param.cell(2, 1).add_paragraph(f'{data["I"]}')
        self.table_param.cell(3, 1).add_paragraph(f'{data["x"]}')
        self.table_param.cell(4, 1).add_paragraph(f'{data["y"]}')
        self.table_param.cell(5, 1).add_paragraph(f'{data["Angle"]}')



    def write_param_data(self):
        data = self.data['Parameters']
        self.table_param.cell(0, 0).add_paragraph(f'L(cm)')
        self.table_param.cell(1, 0).add_paragraph(f'R(cm)')
        self.table_param.cell(2, 0).add_paragraph(f'I(A)')
    def write_meas_data(self):
        data = self.data['Measurements']['Harm']
        self.table_meas.cell(0, 0).add_paragraph(f'Harmonic')
        self.table_meas.cell(0, 1).add_paragraph(f'A(Gs)')
        self.table_meas.cell(0, 2).add_paragraph(f'B(Gs)')
        self.table_meas.cell(0, 3).add_paragraph(f'Amplitude(Gs)')
        self.table_meas.cell(0, 4).add_paragraph(f'Relative')
        for i in range(1, 15):
            self.table_meas.cell(i, 0).add_paragraph(f'{i}')
            self.table_meas.cell(i, 1).add_paragraph(f'{data[i]["A"]}')
            self.table_meas.cell(i, 2).add_paragraph(f'{data[i]["B"]}')
            self.table_meas.cell(i, 3).add_paragraph(f'{data[i]["Amp"]}')
            self.table_meas.cell(i, 4).add_paragraph(f'{data[i]["Rel"]}')
